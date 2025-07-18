"""
Content Analyzer

Analyzes file contents using AI and specialized parsers to generate
comprehensive descriptions including metadata, format analysis, 
and processing recommendations.
"""

import os
import json
import asyncio
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from datetime import datetime
from dataclasses import dataclass, asdict

# Standard library parsers
import csv
import xml.etree.ElementTree as ET
from urllib.parse import urlparse

# Optional parsers
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    from ydata_profiling import ProfileReport
    HAS_YDATA_PROFILING = True
except ImportError:
    HAS_YDATA_PROFILING = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from zohar.tools.camel_tool_manager import CamelToolManager
from zohar.utils.logging import get_logger
from .format_detector import FormatInfo

logger = get_logger(__name__)


@dataclass
class ContentDescription:
    """Comprehensive content description."""
    file_path: str
    format_info: FormatInfo
    
    # Basic properties
    is_human_readable: bool
    is_machine_readable: bool
    
    # File attributes
    file_attributes: Dict[str, Any]
    
    # Metadata extraction
    metadata: Dict[str, Any]
    
    # Content analysis
    content_analysis: Dict[str, Any]
    
    # Processing recommendations
    processing_recommendations: List[str]
    
    # AI-generated description
    ai_description: str
    
    # Parsing results
    parsing_results: Dict[str, Any]
    
    # Quality assessment
    quality_score: float
    quality_issues: List[str]


class ContentAnalyzer:
    """
    Analyzes file contents using multiple approaches:
    1. Specialized parsers for known formats
    2. AI-powered content analysis
    3. Metadata extraction
    4. Quality assessment
    5. Processing recommendations
    """
    
    def __init__(self):
        self.tool_manager = CamelToolManager()
        
        # Parser registry
        self.parsers = {
            'csv': self._parse_csv,
            'json': self._parse_json,
            'xml': self._parse_xml,
            'html': self._parse_html,
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'text': self._parse_text,
            'log': self._parse_log,
            'yaml': self._parse_yaml,
            'markdown': self._parse_markdown
        }
        
        # Analysis prompts for AI
        self.analysis_prompts = {
            'description': """
            Analyze this file content and provide a comprehensive description including:
            1. What type of data this contains
            2. Key characteristics and structure
            3. Potential use cases
            4. Any notable patterns or metadata
            5. Data quality observations
            
            Content preview: {content_preview}
            File info: {file_info}
            """,
            
            'metadata_extraction': """
            Extract structured metadata from this content including:
            - Time/date information
            - Location data
            - People/entities mentioned
            - Categories or tags
            - Version information
            - Source information
            
            Content: {content}
            """,
            
            'recommendations': """
            Based on this file analysis, provide specific recommendations for:
            1. How to process this file
            2. What parsers or tools to use
            3. Potential data quality issues to address
            4. Optimal data structure for vectorization
            
            Analysis: {analysis}
            Format: {format_info}
            """
        }
    
    async def analyze_content(self, file_path: str, format_info: FormatInfo, 
                            max_content_length: int = 10000) -> ContentDescription:
        """
        Perform comprehensive content analysis.
        
        Args:
            file_path: Path to the file
            format_info: Format detection results
            max_content_length: Maximum content length to analyze
            
        Returns:
            ContentDescription with analysis results
        """
        logger.info(f"Analyzing content of: {file_path}")
        
        # Extract file attributes
        file_attributes = self._extract_file_attributes(file_path)
        
        # Parse content using appropriate parser
        parsing_results = await self._parse_content(file_path, format_info, max_content_length)
        
        # Determine readability
        is_human_readable, is_machine_readable = self._assess_readability(format_info, parsing_results)
        
        # Extract metadata
        metadata = await self._extract_metadata(file_path, format_info, parsing_results)
        
        # Perform content analysis
        content_analysis = await self._analyze_content_structure(parsing_results, format_info)
        
        # Generate AI description
        ai_description = await self._generate_ai_description(
            file_path, format_info, parsing_results, content_analysis
        )
        
        # Get processing recommendations
        recommendations = await self._generate_recommendations(
            format_info, content_analysis, parsing_results
        )
        
        # Assess quality
        quality_score, quality_issues = self._assess_quality(parsing_results, content_analysis)
        
        return ContentDescription(
            file_path=file_path,
            format_info=format_info,
            is_human_readable=is_human_readable,
            is_machine_readable=is_machine_readable,
            file_attributes=file_attributes,
            metadata=metadata,
            content_analysis=content_analysis,
            processing_recommendations=recommendations,
            ai_description=ai_description,
            parsing_results=parsing_results,
            quality_score=quality_score,
            quality_issues=quality_issues
        )
    
    def _extract_file_attributes(self, file_path: str) -> Dict[str, Any]:
        """Extract basic file attributes."""
        try:
            path = Path(file_path)
            stat = path.stat()
            
            return {
                'size_bytes': stat.st_size,
                'size_human': self._human_readable_size(stat.st_size),
                'created_timestamp': stat.st_ctime,
                'modified_timestamp': stat.st_mtime,
                'created_date': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_date': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'permissions': oct(stat.st_mode),
                'is_symlink': path.is_symlink(),
                'absolute_path': str(path.absolute()),
                'parent_directory': str(path.parent),
                'filename': path.name,
                'stem': path.stem,
                'extension': path.suffix
            }
        except Exception as e:
            logger.error(f"Failed to extract file attributes for {file_path}: {e}")
            return {}
    
    async def _parse_content(self, file_path: str, format_info: FormatInfo, 
                           max_length: int) -> Dict[str, Any]:
        """Parse content using appropriate parser."""
        detected_format = format_info.detected_format
        
        # Get appropriate parser
        parser = self.parsers.get(detected_format, self._parse_generic)
        
        try:
            return await parser(file_path, format_info, max_length)
        except Exception as e:
            logger.error(f"Failed to parse {file_path} as {detected_format}: {e}")
            # Fallback to generic parsing
            return await self._parse_generic(file_path, format_info, max_length)
    
    async def _parse_csv(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse CSV file."""
        results = {
            'parser': 'csv',
            'success': False,
            'error': None,
            'preview': None,
            'structure': None,
            'statistics': None
        }
        
        try:
            encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
            
            # Basic CSV parsing
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                # Read first few lines for preview
                preview_lines = []
                for i, line in enumerate(f):
                    if i >= 10:  # Limit preview
                        break
                    preview_lines.append(line.strip())
                
                results['preview'] = preview_lines
            
            # Try pandas for better analysis if available
            if HAS_PANDAS:
                try:
                    df = pd.read_csv(file_path, encoding=encoding, nrows=1000)  # Limit rows
                    
                    results['structure'] = {
                        'columns': list(df.columns),
                        'column_count': len(df.columns),
                        'row_count': len(df),
                        'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()},
                        'null_counts': df.isnull().sum().to_dict(),
                        'sample_values': {col: df[col].dropna().head(3).tolist() 
                                        for col in df.columns}
                    }
                    
                    # Generate profile report if ydata-profiling is available
                    if HAS_YDATA_PROFILING and len(df) > 0:
                        try:
                            profile = ProfileReport(df, minimal=True)
                            results['statistics'] = {
                                'profile_generated': True,
                                'warnings': len(profile.get_description()['alerts']),
                                'missing_data': profile.get_description()['table']['n_cells_missing'],
                                'duplicate_rows': profile.get_description()['table']['n_duplicates']
                            }
                        except Exception as e:
                            logger.debug(f"Failed to generate profile for {file_path}: {e}")
                    
                except Exception as e:
                    logger.debug(f"Pandas parsing failed for {file_path}: {e}")
            
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"CSV parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_json(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse JSON file."""
        results = {
            'parser': 'json',
            'success': False,
            'error': None,
            'structure': None,
            'preview': None
        }
        
        try:
            encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read(max_length)
                
            data = json.loads(content)
            
            results['structure'] = {
                'type': type(data).__name__,
                'size': len(data) if isinstance(data, (list, dict)) else 1,
                'keys': list(data.keys()) if isinstance(data, dict) else None,
                'first_few_items': data[:3] if isinstance(data, list) else None,
                'sample_structure': self._analyze_json_structure(data)
            }
            
            results['preview'] = content[:500]  # First 500 chars
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"JSON parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_xml(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse XML file."""
        results = {
            'parser': 'xml',
            'success': False,
            'error': None,
            'structure': None,
            'preview': None
        }
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            results['structure'] = {
                'root_tag': root.tag,
                'namespaces': dict(root.nsmap) if hasattr(root, 'nsmap') else {},
                'child_count': len(root),
                'attributes': root.attrib,
                'child_tags': list(set(child.tag for child in root)),
                'depth': self._calculate_xml_depth(root)
            }
            
            # Read preview
            encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                results['preview'] = f.read(500)
            
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"XML parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_html(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse HTML file."""
        results = {
            'parser': 'html',
            'success': False,
            'error': None,
            'structure': None,
            'preview': None
        }
        
        try:
            encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read(max_length)
            
            if HAS_BS4:
                soup = BeautifulSoup(content, 'html.parser')
                
                results['structure'] = {
                    'title': soup.title.string if soup.title else None,
                    'meta_tags': [tag.attrs for tag in soup.find_all('meta')],
                    'links': [link.get('href') for link in soup.find_all('a', href=True)][:10],
                    'images': [img.get('src') for img in soup.find_all('img', src=True)][:10],
                    'scripts': [script.get('src') for script in soup.find_all('script', src=True)][:10],
                    'text_content_length': len(soup.get_text()),
                    'tag_counts': self._count_html_tags(soup)
                }
            else:
                # Basic HTML analysis without BeautifulSoup
                results['structure'] = {
                    'has_doctype': content.lower().startswith('<!doctype'),
                    'has_html_tag': '<html' in content.lower(),
                    'has_head': '<head>' in content.lower(),
                    'has_body': '<body>' in content.lower(),
                    'content_length': len(content)
                }
            
            results['preview'] = content[:500]
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"HTML parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_pdf(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse PDF file."""
        results = {
            'parser': 'pdf',
            'success': False,
            'error': None,
            'structure': None,
            'text_content': None
        }
        
        if not HAS_PYPDF2:
            results['error'] = "PyPDF2 not available"
            return results
        
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                results['structure'] = {
                    'page_count': len(reader.pages),
                    'metadata': dict(reader.metadata) if reader.metadata else {},
                    'encrypted': reader.is_encrypted,
                    'form_fields': bool(reader.form) if hasattr(reader, 'form') else False
                }
                
                # Extract text from first few pages
                text_content = []
                for i, page in enumerate(reader.pages[:3]):  # First 3 pages
                    try:
                        text_content.append(page.extract_text())
                    except Exception:
                        text_content.append("[Text extraction failed]")
                
                results['text_content'] = text_content
                results['success'] = True
                
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"PDF parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_docx(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse DOCX file."""
        results = {
            'parser': 'docx',
            'success': False,
            'error': None,
            'structure': None,
            'text_content': None
        }
        
        if not HAS_DOCX:
            results['error'] = "python-docx not available"
            return results
        
        try:
            doc = docx.Document(file_path)
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            results['structure'] = {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'section_count': len(doc.sections),
                'word_count': sum(len(p.split()) for p in paragraphs),
                'has_images': len(doc.inline_shapes) > 0
            }
            
            results['text_content'] = paragraphs[:5]  # First 5 paragraphs
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"DOCX parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_text(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse plain text file."""
        results = {
            'parser': 'text',
            'success': False,
            'error': None,
            'content': None,
            'statistics': None
        }
        
        try:
            encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read(max_length)
            
            lines = content.split('\n')
            words = content.split()
            
            results['content'] = content
            results['statistics'] = {
                'character_count': len(content),
                'line_count': len(lines),
                'word_count': len(words),
                'average_line_length': sum(len(line) for line in lines) / len(lines) if lines else 0,
                'empty_lines': sum(1 for line in lines if not line.strip()),
                'longest_line_length': max(len(line) for line in lines) if lines else 0
            }
            
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"Text parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_log(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse log file."""
        results = {
            'parser': 'log',
            'success': False,
            'error': None,
            'log_analysis': None,
            'sample_entries': None
        }
        
        try:
            encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                lines = []
                for i, line in enumerate(f):
                    if i >= 100:  # First 100 lines
                        break
                    lines.append(line.strip())
            
            # Analyze log patterns
            level_counts = {}
            timestamp_patterns = []
            
            for line in lines:
                # Count log levels
                line_lower = line.lower()
                for level in ['error', 'warn', 'info', 'debug', 'trace', 'fatal']:
                    if level in line_lower:
                        level_counts[level] = level_counts.get(level, 0) + 1
                
                # Look for timestamp patterns
                if any(char.isdigit() for char in line[:20]):
                    timestamp_patterns.append(line[:20])
            
            results['log_analysis'] = {
                'total_lines_analyzed': len(lines),
                'log_level_distribution': level_counts,
                'has_timestamps': len(timestamp_patterns) > 0,
                'sample_timestamps': timestamp_patterns[:5],
                'average_line_length': sum(len(line) for line in lines) / len(lines) if lines else 0
            }
            
            results['sample_entries'] = lines[:10]
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"Log parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_yaml(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse YAML file."""
        results = {
            'parser': 'yaml',
            'success': False,
            'error': None,
            'structure': None,
            'content': None
        }
        
        try:
            import yaml
            
            encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read(max_length)
            
            data = yaml.safe_load(content)
            
            results['structure'] = {
                'type': type(data).__name__,
                'keys': list(data.keys()) if isinstance(data, dict) else None,
                'size': len(data) if isinstance(data, (list, dict)) else 1,
                'nested_levels': self._calculate_yaml_depth(data)
            }
            
            results['content'] = content[:500]
            results['success'] = True
            
        except ImportError:
            results['error'] = "PyYAML not available"
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"YAML parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_markdown(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Parse Markdown file."""
        results = {
            'parser': 'markdown',
            'success': False,
            'error': None,
            'structure': None,
            'content': None
        }
        
        try:
            encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read(max_length)
            
            lines = content.split('\n')
            
            # Analyze markdown structure
            headers = []
            lists = 0
            code_blocks = 0
            links = 0
            images = 0
            
            for line in lines:
                line = line.strip()
                if line.startswith('#'):
                    headers.append(line)
                elif line.startswith(('- ', '* ', '+ ')) or line.startswith(tuple(f'{i}. ' for i in range(10))):
                    lists += 1
                elif line.startswith('```'):
                    code_blocks += 1
                elif '[' in line and '](' in line:
                    if line.count('[') > line.count('!['):
                        links += line.count('[') - line.count('![')
                    images += line.count('![')
            
            results['structure'] = {
                'headers': headers,
                'header_count': len(headers),
                'list_items': lists,
                'code_blocks': code_blocks // 2,  # Pairs of ```
                'links': links,
                'images': images,
                'line_count': len(lines),
                'word_count': len(content.split())
            }
            
            results['content'] = content
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"Markdown parsing failed for {file_path}: {e}")
        
        return results
    
    async def _parse_generic(self, file_path: str, format_info: FormatInfo, max_length: int) -> Dict[str, Any]:
        """Generic fallback parser."""
        results = {
            'parser': 'generic',
            'success': False,
            'error': None,
            'content_preview': None,
            'basic_analysis': None
        }
        
        try:
            if format_info.is_text:
                encoding = format_info.encoding if format_info.encoding != 'unknown' else 'utf-8'
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read(max_length)
                    
                results['content_preview'] = content[:500]
                results['basic_analysis'] = {
                    'character_count': len(content),
                    'line_count': len(content.split('\n')),
                    'word_count': len(content.split()),
                    'contains_urls': 'http' in content.lower(),
                    'contains_emails': '@' in content and '.' in content
                }
            else:
                # Binary file - just read bytes
                with open(file_path, 'rb') as f:
                    bytes_content = f.read(min(max_length, 1024))
                    
                results['content_preview'] = f"Binary content: {len(bytes_content)} bytes"
                results['basic_analysis'] = {
                    'size_bytes': len(bytes_content),
                    'null_bytes': bytes_content.count(b'\x00'),
                    'printable_ratio': sum(1 for b in bytes_content if 32 <= b <= 126) / len(bytes_content) if bytes_content else 0
                }
            
            results['success'] = True
            
        except Exception as e:
            results['error'] = str(e)
            logger.error(f"Generic parsing failed for {file_path}: {e}")
        
        return results
    
    def _assess_readability(self, format_info: FormatInfo, parsing_results: Dict) -> tuple[bool, bool]:
        """Assess if file is human readable and/or machine readable."""
        is_human_readable = format_info.is_text and parsing_results.get('success', False)
        
        machine_readable_formats = {
            'csv', 'json', 'xml', 'yaml', 'html', 'sql', 'config'
        }
        is_machine_readable = format_info.detected_format in machine_readable_formats
        
        return is_human_readable, is_machine_readable
    
    async def _extract_metadata(self, file_path: str, format_info: FormatInfo, 
                              parsing_results: Dict) -> Dict[str, Any]:
        """Extract metadata from file content."""
        metadata = {
            'timestamps': [],
            'locations': [],
            'entities': [],
            'categories': [],
            'version_info': {},
            'source_info': {}
        }
        
        # Extract from parsing results first
        if parsing_results.get('success') and 'structure' in parsing_results:
            structure = parsing_results['structure']
            
            # Extract metadata from structure
            if isinstance(structure, dict):
                if 'metadata' in structure:
                    metadata['source_info'].update(structure['metadata'])
                if 'title' in structure:
                    metadata['source_info']['title'] = structure['title']
        
        # Use AI for advanced metadata extraction if content is available
        content_for_ai = self._get_content_for_ai(parsing_results)
        if content_for_ai:
            try:
                ai_metadata = await self._extract_ai_metadata(content_for_ai)
                if ai_metadata:
                    metadata.update(ai_metadata)
            except Exception as e:
                logger.debug(f"AI metadata extraction failed: {e}")
        
        return metadata
    
    async def _analyze_content_structure(self, parsing_results: Dict, 
                                       format_info: FormatInfo) -> Dict[str, Any]:
        """Analyze content structure and patterns."""
        analysis = {
            'complexity_score': 0.0,
            'structure_type': 'unknown',
            'data_patterns': [],
            'content_quality': 'unknown',
            'suggested_schema': None
        }
        
        if not parsing_results.get('success'):
            return analysis
        
        # Analyze based on format
        if format_info.detected_format == 'csv':
            analysis.update(self._analyze_csv_structure(parsing_results))
        elif format_info.detected_format == 'json':
            analysis.update(self._analyze_json_structure_detailed(parsing_results))
        elif format_info.detected_format in ['xml', 'html']:
            analysis.update(self._analyze_markup_structure(parsing_results))
        elif format_info.detected_format == 'text':
            analysis.update(self._analyze_text_structure(parsing_results))
        
        return analysis
    
    async def _generate_ai_description(self, file_path: str, format_info: FormatInfo,
                                     parsing_results: Dict, content_analysis: Dict) -> str:
        """Generate AI-powered content description."""
        try:
            content_preview = self._get_content_for_ai(parsing_results, max_length=1000)
            
            prompt = self.analysis_prompts['description'].format(
                content_preview=content_preview,
                file_info=f"Format: {format_info.detected_format}, Size: {format_info.file_path}"
            )
            
            # Use CAMEL AI for description generation
            response = await self.tool_manager.execute_with_agent(
                prompt, toolkit_filter=['search']  # Only use search tools if needed
            )
            
            return response.get('content', 'AI description generation failed')
            
        except Exception as e:
            logger.error(f"AI description generation failed: {e}")
            return f"Content analysis completed. Format: {format_info.detected_format}. " \
                   f"File appears to be {'human readable' if format_info.is_text else 'binary data'}."
    
    async def _generate_recommendations(self, format_info: FormatInfo, 
                                      content_analysis: Dict, parsing_results: Dict) -> List[str]:
        """Generate processing recommendations."""
        recommendations = []
        
        # Format-specific recommendations
        format_type = format_info.detected_format
        
        if format_type == 'csv':
            recommendations.extend([
                "Use pandas.read_csv() for data analysis",
                "Consider data type inference and null value handling",
                "Check for delimiter consistency"
            ])
        elif format_type == 'json':
            recommendations.extend([
                "Use json.loads() or pandas.read_json() for parsing",
                "Validate JSON schema if structured data",
                "Consider flattening nested structures for analysis"
            ])
        elif format_type == 'pdf':
            recommendations.extend([
                "Use PyPDF2 or pdfplumber for text extraction",
                "Consider OCR for scanned documents",
                "Extract metadata and structure information"
            ])
        elif format_info.is_text:
            recommendations.extend([
                "Verify text encoding before processing",
                "Consider chunking for large files",
                "Apply appropriate text preprocessing"
            ])
        
        # Quality-based recommendations
        if parsing_results.get('error'):
            recommendations.append("Address parsing errors before processing")
        
        # Add encoding recommendations
        if format_info.encoding == 'unknown':
            recommendations.append("Detect and verify text encoding")
        
        return recommendations[:10]  # Limit to 10 recommendations
    
    def _assess_quality(self, parsing_results: Dict, content_analysis: Dict) -> tuple[float, List[str]]:
        """Assess content quality and identify issues."""
        score = 1.0
        issues = []
        
        # Parsing success
        if not parsing_results.get('success'):
            score -= 0.3
            issues.append("Failed to parse content")
        
        # Format-specific quality checks
        if 'structure' in parsing_results:
            structure = parsing_results['structure']
            
            # CSV quality
            if 'null_counts' in structure:
                null_ratio = sum(structure['null_counts'].values()) / max(1, sum(structure['null_counts'].values()))
                if null_ratio > 0.3:
                    score -= 0.2
                    issues.append("High missing data ratio")
            
            # General structure issues
            if isinstance(structure, dict) and not structure:
                score -= 0.1
                issues.append("Empty or minimal structure")
        
        # Content analysis quality
        complexity = content_analysis.get('complexity_score', 0)
        if complexity < 0.3:
            issues.append("Low content complexity")
        elif complexity > 0.9:
            issues.append("Very high complexity may need special handling")
        
        return max(0.0, score), issues
    
    # Helper methods
    def _human_readable_size(self, size_bytes: int) -> str:
        """Convert bytes to human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f}{unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f}PB"
    
    def _analyze_json_structure(self, data: Any, max_depth: int = 3) -> Dict[str, Any]:
        """Analyze JSON structure recursively."""
        if max_depth <= 0:
            return {'truncated': True}
        
        if isinstance(data, dict):
            return {
                'type': 'object',
                'keys': list(data.keys())[:10],  # First 10 keys
                'key_count': len(data),
                'nested_objects': sum(1 for v in data.values() if isinstance(v, dict)),
                'nested_arrays': sum(1 for v in data.values() if isinstance(v, list))
            }
        elif isinstance(data, list):
            return {
                'type': 'array',
                'length': len(data),
                'item_types': list(set(type(item).__name__ for item in data[:10])),
                'first_item_structure': self._analyze_json_structure(data[0], max_depth - 1) if data else None
            }
        else:
            return {
                'type': type(data).__name__,
                'value': str(data)[:100] if len(str(data)) <= 100 else str(data)[:100] + '...'
            }
    
    def _calculate_xml_depth(self, element, current_depth: int = 0) -> int:
        """Calculate maximum depth of XML tree."""
        if not element:
            return current_depth
        return max([self._calculate_xml_depth(child, current_depth + 1) 
                   for child in element] + [current_depth])
    
    def _count_html_tags(self, soup) -> Dict[str, int]:
        """Count occurrences of HTML tags."""
        tag_counts = {}
        for tag in soup.find_all():
            tag_name = tag.name
            tag_counts[tag_name] = tag_counts.get(tag_name, 0) + 1
        return dict(sorted(tag_counts.items(), key=lambda x: x[1], reverse=True)[:10])
    
    def _calculate_yaml_depth(self, data: Any, current_depth: int = 0) -> int:
        """Calculate maximum depth of YAML structure."""
        if isinstance(data, dict):
            if not data:
                return current_depth
            return max(self._calculate_yaml_depth(v, current_depth + 1) for v in data.values())
        elif isinstance(data, list):
            if not data:
                return current_depth
            return max(self._calculate_yaml_depth(item, current_depth + 1) for item in data)
        else:
            return current_depth
    
    def _get_content_for_ai(self, parsing_results: Dict, max_length: int = 500) -> str:
        """Extract content suitable for AI analysis."""
        if not parsing_results.get('success'):
            return ""
        
        # Try different content fields
        for field in ['content', 'preview', 'text_content', 'content_preview']:
            if field in parsing_results and parsing_results[field]:
                content = parsing_results[field]
                if isinstance(content, list):
                    content = '\n'.join(str(item) for item in content)
                elif not isinstance(content, str):
                    content = str(content)
                
                return content[:max_length]
        
        return ""
    
    async def _extract_ai_metadata(self, content: str) -> Dict[str, Any]:
        """Use AI to extract metadata from content."""
        # This is a placeholder for AI-powered metadata extraction
        # In a real implementation, you would use CAMEL AI tools here
        return {
            'ai_extracted': True,
            'confidence': 0.5
        }
    
    def _analyze_csv_structure(self, parsing_results: Dict) -> Dict[str, Any]:
        """Analyze CSV structure in detail."""
        structure = parsing_results.get('structure', {})
        
        return {
            'complexity_score': min(1.0, structure.get('column_count', 0) / 20),
            'structure_type': 'tabular',
            'data_patterns': [f"Columns: {structure.get('column_count', 0)}"],
            'content_quality': 'good' if structure.get('row_count', 0) > 0 else 'poor'
        }
    
    def _analyze_json_structure_detailed(self, parsing_results: Dict) -> Dict[str, Any]:
        """Analyze JSON structure in detail."""
        structure = parsing_results.get('structure', {})
        
        complexity = 0.5
        if structure.get('type') == 'dict':
            complexity += len(structure.get('keys', [])) / 50
        elif structure.get('type') == 'list':
            complexity += min(1.0, structure.get('size', 0) / 100)
        
        return {
            'complexity_score': min(1.0, complexity),
            'structure_type': structure.get('type', 'unknown'),
            'data_patterns': [f"Type: {structure.get('type', 'unknown')}"]
        }
    
    def _analyze_markup_structure(self, parsing_results: Dict) -> Dict[str, Any]:
        """Analyze XML/HTML structure."""
        structure = parsing_results.get('structure', {})
        
        return {
            'complexity_score': min(1.0, structure.get('child_count', 0) / 50),
            'structure_type': 'hierarchical',
            'data_patterns': [f"Children: {structure.get('child_count', 0)}"],
            'content_quality': 'good' if structure.get('child_count', 0) > 0 else 'poor'
        }
    
    def _analyze_text_structure(self, parsing_results: Dict) -> Dict[str, Any]:
        """Analyze plain text structure."""
        stats = parsing_results.get('statistics', {})
        
        return {
            'complexity_score': min(1.0, stats.get('word_count', 0) / 1000),
            'structure_type': 'unstructured',
            'data_patterns': [f"Lines: {stats.get('line_count', 0)}",
                            f"Words: {stats.get('word_count', 0)}"],
            'content_quality': 'good' if stats.get('word_count', 0) > 0 else 'poor'
        } 